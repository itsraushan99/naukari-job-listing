import re
import io
import PyPDF2
import docx
from pathlib import Path


def extract_text_from_pdf(source):
    """Extract text from PDF — accepts file path or bytes"""
    text = ""
    try:
        if isinstance(source, (str, Path)):
            with open(source, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text()
        else:
            reader = PyPDF2.PdfReader(io.BytesIO(source))
            for page in reader.pages:
                text += page.extract_text()
    except:
        pass
    return text


def extract_text_from_docx(source):
    """Extract text from DOCX — accepts file path or bytes"""
    text = ""
    try:
        if isinstance(source, (str, Path)):
            doc = docx.Document(source)
        else:
            doc = docx.Document(io.BytesIO(source))
        for para in doc.paragraphs:
            text += para.text + "\n"
    except:
        pass
    return text


def extract_text_from_txt(source):
    """Extract text from TXT — accepts file path or bytes"""
    try:
        if isinstance(source, (str, Path)):
            with open(source, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            return source.decode('utf-8')
    except:
        return ""


def _parse_text(text):
    """Core parsing logic — shared by both parse functions"""
    if not text:
        return None
    text_lower = text.lower()
    return {
        'experience': extract_experience(text_lower),
        'job_titles': extract_job_titles(text_lower),
        'location': extract_location(text_lower),
        'company_types': extract_company_preference(text_lower),
        'raw_text': text
    }


def parse_resume(file_path):
    """Parse resume from a file path on disk"""
    file_path = Path(file_path)
    ext = file_path.suffix.lower()
    if ext == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        text = extract_text_from_docx(file_path)
    elif ext == '.txt':
        text = extract_text_from_txt(file_path)
    else:
        return None
    return _parse_text(text)


def parse_resume_from_bytes(file_bytes, filename):
    """Parse resume from raw bytes (used for st.file_uploader)"""
    ext = Path(filename).suffix.lower()
    if ext == '.pdf':
        text = extract_text_from_pdf(file_bytes)
    elif ext in ['.docx', '.doc']:
        text = extract_text_from_docx(file_bytes)
    elif ext == '.txt':
        text = extract_text_from_txt(file_bytes)
    else:
        return None
    return _parse_text(text)


def extract_experience(text):
    patterns = [
        r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
        r'experience[:\s]+(\d+)\+?\s*years?',
        r'(\d+)\+?\s*yrs?\s+(?:of\s+)?experience'
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return int(match.group(1))
    return 0


def extract_job_titles(text):
    job_mapping = {
        'Web Developer': ['html', 'css', 'javascript', 'web development', 'frontend', 'backend'],
        'Python Developer': ['python', 'django', 'flask', 'fastapi'],
        'Data Analyst': ['data analysis', 'sql', 'excel', 'tableau', 'power bi', 'analytics'],
        'Data Scientist': ['machine learning', 'data science', 'ml', 'ai', 'deep learning'],
        'Full Stack Developer': ['full stack', 'mern', 'mean', 'fullstack'],
        'Frontend Developer': ['react', 'angular', 'vue', 'frontend', 'ui'],
        'Backend Developer': ['backend', 'api', 'rest', 'node', 'express'],
        'DevOps Engineer': ['devops', 'docker', 'kubernetes', 'ci/cd', 'jenkins', 'aws'],
        'Java Developer': ['java', 'spring', 'hibernate', 'j2ee'],
        'React Developer': ['react', 'reactjs', 'redux'],
        'Node.js Developer': ['node', 'nodejs', 'express'],
        'Software Engineer': ['software engineer', 'software development', 'programming'],
        'Machine Learning Engineer': ['machine learning', 'ml engineer', 'tensorflow', 'pytorch'],
        'Cloud Engineer': ['aws', 'azure', 'gcp', 'cloud'],
        'QA Engineer': ['testing', 'qa', 'selenium', 'automation testing'],
        'Business Analyst': ['business analysis', 'requirements', 'stakeholder'],
        'Product Manager': ['product management', 'product manager', 'roadmap'],
        'UI/UX Designer': ['ui', 'ux', 'figma', 'design', 'user experience']
    }
    matched = [title for title, keywords in job_mapping.items() if any(k in text for k in keywords)]
    return matched if matched else ['Web Developer']


def extract_location(text):
    indian_cities = ['bangalore', 'mumbai', 'delhi', 'hyderabad', 'pune', 'chennai',
                     'kolkata', 'ahmedabad', 'noida', 'gurgaon', 'gurugram']
    for city in indian_cities:
        if city in text:
            return city.title()
    return 'India'


def extract_company_preference(text):
    preferences = []
    if any(w in text for w in ['product based', 'product company', 'saas']):
        preferences.append('Product Based')
    if any(w in text for w in ['service based', 'consulting', 'it services']):
        preferences.append('Service Based')
    if any(w in text for w in ['startup', 'early stage']):
        preferences.append('Startup')
    return preferences
